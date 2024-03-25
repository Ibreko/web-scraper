from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from selenium.common.exceptions import NoSuchElementException

# Create a new word document
doc = Document()


desired_capabilities = {
    "acceptInsecureCerts": True
}

driver = webdriver.Chrome()
driver.get('https://vulmon.com/')

input="CVE-2017-0145"
search_box = driver.find_element(by='id', value='inputSearch')
search_query = input
search_box.send_keys(search_query)
search_box.send_keys(Keys.RETURN)


wait = WebDriverWait(driver, 100000)
wait.until(EC.presence_of_element_located((By.CLASS_NAME, 'content')))
result_elements = driver.find_elements(By.CLASS_NAME, 'item')

link_text_to_click = input
link_to_click = driver.find_element(By.LINK_TEXT, link_text_to_click)

link_to_click.click()

try:
     element1 = driver.find_element(By.XPATH, '//*[@id="cust-css-overrides"]/div[3]/div/div[1]/div[1]/div/div')
     data1 = element1.text
     print(data1)

except NoSuchElementException:
    data1 = "XPath not found"
    print("XPath not found for element2")
try:
    element2 = driver.find_element(By.XPATH, '//*[@id="cust-css-overrides"]/div[3]/div/div[1]/div[2]')
    data2 = element2.text
    print(data2)
except NoSuchElementException:
    data2 = "XPath not found"
    print("XPath not found for element2")

try:
    element3 = driver.find_element(By.XPATH, '//*[@id="cust-css-overrides"]/div[3]/div/div[1]/div[7]')
    data3 = element3.text
    print(data3)
except NoSuchElementException:
    data3 = "XPath not found"
    print("XPath not found for element3")

try: 
     element4 = driver.find_element(By.XPATH, '//*[@id="cust-css-overrides"]/div[3]/div/div[1]/div[8]')
     data4 = element4.text
     print(data4)

except NoSuchElementException:
    data4 = "XPath not found"
    print("XPath not found for element1")

try:
     element5 = driver.find_element(By.XPATH,'//*[@id="cust-css-overrides"]/div[3]/div/div[1]/div[9]')
     data5=element5.text
     print(data5)

except NoSuchElementException:
    data5 = "XPath not found"
    print("XPath not found for element1")

try:
     element6 = driver.find_element(By.XPATH,'//*[@id="cust-css-overrides"]/div[3]/div/div[1]/div[10]')
     data6=element6.text
     print(data6)

except NoSuchElementException:
    data6 = "XPath not found"
    print("XPath not found for element1")

try:
     element7 = driver.find_element(By.XPATH,'//*[@id="cust-css-overrides"]/div[3]/div/div[1]/div[11]')
     data7=element7.text
     print(data7)

except NoSuchElementException:
    data7 = "XPath not found"
    print("XPath not found for element1")

driver.get('https://nvd.nist.gov/vuln/search')


search_box = driver.find_element(by='id', value='Keywords')
search_query = input
search_box.send_keys(search_query)
search_box.send_keys(Keys.RETURN)


wait = WebDriverWait(driver, 100000)
wait.until(EC.presence_of_element_located((By.CLASS_NAME, 'container')))
result_elements = driver.find_elements(By.CLASS_NAME, 'container')

link_text_to_click = input
link_to_click = driver.find_element(By.LINK_TEXT, link_text_to_click)


link_to_click.click()

try:
      element8 = driver.find_element(By.XPATH,'//*[@id="vulnDetailTableView"]/tbody/tr/td/h2')
      data8=element8.text
      print(data8)

except NoSuchElementException:
    data8 = "XPath not found"
    print("XPath not found for element1")

try:
      element9 = driver.find_element(By.XPATH,'//*[@id="vulnShowWarningDiv"]/div')
      data9=element9.text
      print(data9)

except NoSuchElementException:
    data9 = "XPath not found"
    print("XPath not found for element1")

try:
     element10 = driver.find_element(By.XPATH,'//*[@id="vulnDescriptionTitle"]')
     data10=element10.text
     print(data10)

except NoSuchElementException:
    data10 = "XPath not found"
    print("XPath not found for element1")

try:
     element11 = driver.find_element(By.XPATH,'//*[@id="vulnDetailTableView"]/tbody/tr/td/div/div[1]/p')
     data11=element11.text
     print(data11)

except NoSuchElementException:
    data11 = "XPath not found"
    print("XPath not found for element1")

try:
     element12= driver.find_element(By.XPATH,'//*[@id="vulnCvssPanel"]')
     data12=element12.text
     print(data12)

except NoSuchElementException:
    data12 = "XPath not found"
    print("XPath not found for element1")

try:
     element13 = driver.find_element(By.XPATH,'//*[@id="vulnHyperlinksPanel"]/h3')
     data13=element13.text
     print(data13)

except NoSuchElementException:
    data13 = "XPath not found"
    print("XPath not found for element1")

try:
     element14 = driver.find_element(By.XPATH,'//*[@id="vulnHyperlinksPanel"]/p')
     data14=element14.text
     print(data14)
except NoSuchElementException:
    data14 = "XPath not found"
    print("XPath not found for element1")

try: 
     element15 = driver.find_element(By.XPATH,'//*[@id="vulnHyperlinksPanel"]/table')
     data15=element15.text
     print(data15)

except NoSuchElementException:
    data15 = "XPath not found"
    print("XPath not found for element1")


try:
     element16 = driver.find_element(By.XPATH,'//*[@id="vulnCisaExploit"]')
     data16=element16.text
     print(data16)


except NoSuchElementException:
    data16 = "XPath not found"
    print("XPath not found for element1")

try:
     element17 = driver.find_element(By.XPATH,'//*[@id="vulnTechnicalDetailsDiv"]')
     data17=element17.text
     print(data17)

except NoSuchElementException:
     data17 = "XPath not found"
     print("XPath not found for element1")

try:
     element18 = driver.find_element(By.XPATH,'//*[@id="vulnDetailTableView"]/tbody/tr/td/div/div[1]/div[3]/div[4]')
     data18=element18.text
     print(data18)

except NoSuchElementException:
    data18 = "XPath not found"
    print("XPath not found for element1")
driver.get('https://www.exploit-db.com/search?')

search_box = driver.find_element(by='id', value='cveSearch')
search_query = input
search_box.send_keys(search_query)
search_box.send_keys(Keys.RETURN)

try:
    a_element = (EC.element_to_be_clickable((By.CSS_SELECTOR, '#exploits-table > tbody > tr.odd:nth-child(1) > td.text-center:nth-child(2) > a')))
    a_element.click()
except:
    print("Empty bro")



doc.add_heading(input, level=1)
doc.add_paragraph(data1)
doc.add_heading("Vulnerability Summary", level=1)
doc.add_paragraph(data2)
doc.add_paragraph(data3)
doc.add_heading("Metasploit Modules", level=1)
doc.add_paragraph(data4)
doc.add_heading("Github Repositories", level=1)
doc.add_paragraph(data5)
doc.add_heading("Recent Articles", level=1)
doc.add_paragraph(data6)
doc.add_heading("References", level=1)
doc.add_paragraph(data7)
doc.add_heading("Data From nvd.nist.gov", level=1)
doc.add_paragraph(data8)
doc.add_paragraph(data10)
doc.add_paragraph(data9)
doc.add_heading("References to Advisories,Solutions and Tools", level=1)
doc.add_paragraph(data11)
doc.add_paragraph(data12)
doc.add_heading("This CVE is in CISA's Known Exploited Vulnerabilities Catalog", level=1)
doc.add_paragraph(data13)
doc.add_paragraph(data14)
doc.add_paragraph(data15)
doc.add_paragraph(data16)
doc.add_paragraph(data17)
doc.add_paragraph(data18)


doc.save("BEST PROJECT EVER FOR SIR JAMES2!")
driver.quit()